#!/home/ubuntu/telegram-file-uploader/venv/bin/python3
# -*- coding: utf-8 -*-

"""
محول الملفات المتقدم - MoTech Cloud v4.0
يدعم تحويل Markdown إلى صيغ متعددة مع دعم كامل للغة العربية
"""

import sys
import os
import subprocess
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from fpdf import FPDF
import re

# ═══════════════════════════════════════════════════════
# دوال مساعدة
# ═══════════════════════════════════════════════════════

def clean_markdown(text):
    """إزالة تنسيقات Markdown الأساسية"""
    # إزالة الروابط
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # إزالة التنسيقات
    text = text.replace('**', '').replace('*', '').replace('`', '')
    # إزالة رموز العناوين
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    return text

def parse_markdown_structure(content):
    """تحليل بنية Markdown لاستخراج العناوين والفقرات"""
    lines = content.split('\n')
    structure = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # عناوين
        if line.startswith('# '):
            structure.append({'type': 'h1', 'text': line[2:]})
        elif line.startswith('## '):
            structure.append({'type': 'h2', 'text': line[3:]})
        elif line.startswith('### '):
            structure.append({'type': 'h3', 'text': line[4:]})
        # قوائم
        elif line.startswith('- ') or line.startswith('* '):
            structure.append({'type': 'list', 'text': line[2:]})
        elif re.match(r'^\d+\.\s', line):
            structure.append({'type': 'ordered_list', 'text': re.sub(r'^\d+\.\s', '', line)})
        # فقرات عادية
        else:
            structure.append({'type': 'paragraph', 'text': line})
    
    return structure

# ═══════════════════════════════════════════════════════
# تحويل إلى DOCX
# ═══════════════════════════════════════════════════════

def convert_markdown_to_docx(markdown_content, output_path):
    """تحويل Markdown إلى DOCX مع دعم العربية"""
    try:
        # محاولة استخدام pandoc أولاً
        temp_md_path = "/tmp/temp_content.md"
        with open(temp_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        result = subprocess.run(
            ["pandoc", temp_md_path, "-o", output_path, "--standalone"],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0:
            return True
            
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
        pass
    
    # إذا فشل pandoc، استخدام python-docx
    try:
        doc = Document()
        
        # إعداد الخط العربي
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        structure = parse_markdown_structure(markdown_content)
        
        for item in structure:
            if item['type'] == 'h1':
                p = doc.add_heading(item['text'], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif item['type'] == 'h2':
                p = doc.add_heading(item['text'], level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif item['type'] == 'h3':
                p = doc.add_heading(item['text'], level=3)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif item['type'] in ['list', 'ordered_list']:
                p = doc.add_paragraph(item['text'], style='List Bullet' if item['type'] == 'list' else 'List Number')
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p = doc.add_paragraph(item['text'])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"DOCX Error: {e}", file=sys.stderr)
        return False

# ═══════════════════════════════════════════════════════
# تحويل إلى PDF
# ═══════════════════════════════════════════════════════

def convert_markdown_to_pdf(markdown_content, output_path):
    """تحويل Markdown إلى PDF"""
    try:
        # محاولة استخدام pandoc أولاً
        temp_md_path = "/tmp/temp_content.md"
        with open(temp_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        # محاولة التحويل المباشر
        result = subprocess.run(
            ["pandoc", temp_md_path, "-o", output_path, 
             "--pdf-engine=xelatex", "-V", "mainfont=DejaVu Sans"],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode == 0:
            return True
            
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
        pass
    
    # إذا فشل pandoc، محاولة التحويل عبر HTML
    try:
        temp_html = "/tmp/temp_content.html"
        subprocess.run(
            ["pandoc", temp_md_path, "-o", temp_html, "--standalone"],
            capture_output=True,
            text=True,
            timeout=30,
            check=True
        )
        
        # استخدام wkhtmltopdf إذا كان متاحاً
        subprocess.run(
            ["wkhtmltopdf", temp_html, output_path],
            capture_output=True,
            text=True,
            timeout=60,
            check=True
        )
        
        return True
        
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
        pass
    
    # الخيار الاحتياطي: استخدام fpdf2 (محدود للنصوص البسيطة)
    return convert_markdown_to_pdf_fpdf(markdown_content, output_path)

def convert_markdown_to_pdf_fpdf(markdown_content, output_path):
    """تحويل Markdown إلى PDF باستخدام fpdf2 (احتياطي)"""
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # محاولة استخدام خط يدعم العربية
        try:
            pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
            pdf.set_font('DejaVu', size=12)
        except:
            # إذا فشل، استخدام الخط الافتراضي
            pdf.set_font("Arial", size=12)
        
        # تنظيف المحتوى
        clean_content = clean_markdown(markdown_content)
        
        # كتابة المحتوى
        for line in clean_content.split('\n'):
            if line.strip():
                try:
                    pdf.multi_cell(0, 10, txt=line, align='R')
                except:
                    # إذا فشلت الكتابة، تجاهل السطر
                    pass
        
        pdf.output(output_path)
        return True
        
    except Exception as e:
        print(f"PDF Error: {e}", file=sys.stderr)
        return False

# ═══════════════════════════════════════════════════════
# تحويل إلى Excel
# ═══════════════════════════════════════════════════════

def convert_markdown_to_xlsx(markdown_content, output_path):
    """تحويل Markdown إلى Excel مع تنسيق جيد"""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "المحتوى"
        
        # تنسيق العناوين
        header_font = Font(bold=True, size=14, color="FFFFFF")
        header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
        header_alignment = Alignment(horizontal="right", vertical="center")
        
        # تنسيق النص العادي
        normal_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)
        
        structure = parse_markdown_structure(markdown_content)
        
        row = 1
        for item in structure:
            cell = ws.cell(row=row, column=1, value=item['text'])
            cell.alignment = normal_alignment
            
            if item['type'] in ['h1', 'h2', 'h3']:
                cell.font = Font(bold=True, size=14 if item['type'] == 'h1' else 12)
                cell.fill = PatternFill(start_color="E0E7FF", end_color="E0E7FF", fill_type="solid")
            
            row += 1
        
        # ضبط عرض العمود
        ws.column_dimensions['A'].width = 100
        
        wb.save(output_path)
        return True
        
    except Exception as e:
        print(f"Excel Error: {e}", file=sys.stderr)
        return False

# ═══════════════════════════════════════════════════════
# تحويل إلى PowerPoint
# ═══════════════════════════════════════════════════════

def convert_markdown_to_pptx(markdown_content, output_path):
    """تحويل Markdown إلى PowerPoint"""
    try:
        # محاولة استخدام pandoc
        temp_md_path = "/tmp/temp_content.md"
        with open(temp_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        result = subprocess.run(
            ["pandoc", temp_md_path, "-o", output_path, "-t", "pptx"],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0:
            return True
        else:
            print(f"Pandoc PPTX Error: {result.stderr}", file=sys.stderr)
            return False
            
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
        print(f"PPTX Error: {e}", file=sys.stderr)
        return False

# ═══════════════════════════════════════════════════════
# الدالة الرئيسية
# ═══════════════════════════════════════════════════════

def main():
    """الدالة الرئيسية"""
    if len(sys.argv) != 4:
        print("Usage: python3 converter.py <input_md_path> <output_path> <format>", file=sys.stderr)
        sys.exit(1)

    input_md_path = sys.argv[1]
    output_path = sys.argv[2]
    output_format = sys.argv[3].lower()

    try:
        with open(input_md_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()
    except Exception as e:
        print(f"Error reading input file: {e}", file=sys.stderr)
        sys.exit(1)

    success = False
    
    if output_format == 'docx':
        success = convert_markdown_to_docx(markdown_content, output_path)
    elif output_format == 'pdf':
        success = convert_markdown_to_pdf(markdown_content, output_path)
    elif output_format == 'xlsx':
        success = convert_markdown_to_xlsx(markdown_content, output_path)
    elif output_format == 'pptx':
        success = convert_markdown_to_pptx(markdown_content, output_path)
    elif output_format == 'txt':
        try:
            clean_content = clean_markdown(markdown_content)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(clean_content)
            success = True
        except Exception as e:
            print(f"TXT Error: {e}", file=sys.stderr)
            success = False
    elif output_format == 'md':
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            success = True
        except Exception as e:
            print(f"MD Error: {e}", file=sys.stderr)
            success = False
    else:
        print(f"Unsupported format: {output_format}", file=sys.stderr)
        sys.exit(1)

    if success:
        print(f"✅ Conversion to {output_format.upper()} successful")
        sys.exit(0)
    else:
        print(f"❌ Conversion to {output_format.upper()} failed", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
